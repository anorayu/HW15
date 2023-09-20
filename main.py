# ДЗ 15
from docx import Document

# Создание нового документа
doc = Document()
doc.add_paragraph("Hello Python")

# Сохранение документа
doc.save("hello_python.docx")
from docx import Document

# Открытие файла Word
doc = Document("hello_python.docx")

# Получение жирного текста в строку Python
bold_text = ""
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if run.bold:
            bold_text += run.text

# Вывод жирного текста на экран
print(bold_text)
from docx import Document
from docx.shared import Pt

# Создание нового документа
doc = Document()
paragraph = doc.add_paragraph("This is a paragraph.")

# Изменение шрифта и размера шрифта абзаца
font = paragraph.runs[0].font
font.name = "Arial"
font.size = Pt(12)

# Сохранение документа
doc.save("new_word_file.docx")